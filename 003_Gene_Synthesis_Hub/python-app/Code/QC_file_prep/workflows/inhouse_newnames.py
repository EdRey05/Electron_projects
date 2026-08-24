"""
Bio Basic Inc. - Canada
Made by: Eduardo Reyes, Ph.D.
Contact: ed5reyes@outlook

Version: 3.2
Date: Mar 30, 2026
Notes: Fixed issues with hub-passed directories.
"""

import os
import shutil
import zipfile
from typing import List, Tuple
import tempfile
import pandas as pd
from datetime import date

import customtkinter as ctk
from docx import Document
from docx.shared import Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from PIL import Image
from tkinter import messagebox
from enum import Enum, auto
import pythoncom
import win32com.client as win32

class ProcessStatus(Enum):
    SUCCESS_GEL = auto()
    SUCCESS_NO_GEL = auto()
    SUCCESS_FALLBACK_NO_GEL = auto()
    FAIL_NO_GEL = auto()
    FAIL_DOC_NOT_FOUND = auto()
    FAIL_SUBFOLDER_NOT_FOUND = auto()
    FAIL_UNEXPECTED = auto()
    FAIL_IDS_NOT_FOUND = auto()

class InHouseNewNames:
    def __init__(self, app_instance):
        """
        Initialize the workflow for the 'InHouse NewNames' process.
        - app_instance: A reference to the main App to access common elements.
        """
        self.app = app_instance
        self.log = self.app.log

        # --- Constants for this workflow ---
        self.PROCESSED_FOLDERS_DIR_NAME = "2-processed_folders"
        self.OUTPUT_ZIPS_DIR_NAME = "3-output_zips"
        self.REPORT_SUBSTRING = "gene synthesis report"
        self.QC_TABLE_PARAGRAPH_TEXT = "restriction enzyme digestion analysis"
        self.QC_TABLE_INDEX = 1
        self.NO_GEL_PREFIX = 'X'
        self.GEL_PREFIX = 'Z'
        self.SEQ_PREFIX = 'SEQ'
        self.DEFAULT_MAX_IMG_WIDTH = '2.3'
        self.DEFAULT_MAX_IMG_HEIGHT = '2.8'
        self.BBI_BLUE_COLOR = RGBColor(0x00, 0x70, 0xC0)
        self.WIN32_DOC_FORMAT = 0
        self.WIN32_DOCX_FORMAT = 16
        self.DOC_EXTENSION = ".doc"
        self.DOCX_EXTENSION = ".docx"
        self.TEMP_DOCX_SUFFIX = "_temp.docx"
        self.INVALID_FILE_PREFIX = '~'

        # --- Workflow-specific variables ---
        self.max_w = ctk.StringVar(value=self.app.config.get('DEFAULT', 
                                                            'Max_Img_Width_Inches', 
                                                            fallback=self.DEFAULT_MAX_IMG_WIDTH)
                                                            )
        self.max_h = ctk.StringVar(value=self.app.config.get('DEFAULT', 
                                                            'Max_Img_Height_Inches', 
                                                            fallback=self.DEFAULT_MAX_IMG_HEIGHT)
                                                            )
        self.copy_output_var = ctk.BooleanVar(value=False)
        self.delete_intermediate_var = ctk.BooleanVar(value=False)
        self.rename_from_excel_var = ctk.BooleanVar(value=True)
        
        self.shipment_qc_df = None

    def create_ui(self, parent_frame):
        """
        Creates the UI elements specific to this workflow inside the provided parent frame (a tab).
        """
        parent_frame.grid_columnconfigure(0, weight=13) # 65%
        parent_frame.grid_columnconfigure(1, weight=1)  # 5%
        parent_frame.grid_columnconfigure(2, weight=6)  # 30%

        left_panel = ctk.CTkFrame(parent_frame, fg_color="transparent")
        left_panel.grid(row=0, column=0, sticky="nsew", padx=10)

        right_panel = ctk.CTkFrame(parent_frame, fg_color="transparent")
        right_panel.grid(row=0, column=2, sticky="nsew", padx=10)

        # Max dimensions for QC Gel Image
        dimensions_title_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        dimensions_title_frame.pack(pady=(10, 0), anchor="center")
        ctk.CTkLabel(dimensions_title_frame, text="QC Gel Image Dimensions").pack()

        dimensions_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        dimensions_frame.pack(pady=(4, 8), anchor="center")
        ctk.CTkLabel(dimensions_frame, text="Max Width (in):").pack(side="left", padx=(0, 8))
        ctk.CTkEntry(dimensions_frame, textvariable=self.max_w, width=80).pack(side="left", padx=(0, 16))
        ctk.CTkLabel(dimensions_frame, text="Max Height (in):").pack(side="left", padx=(0, 8))
        ctk.CTkEntry(dimensions_frame, textvariable=self.max_h, width=80).pack(side="left", padx=(0, 16))

        # Optional step checkboxes
        ctk.CTkLabel(right_panel, text="Optional Steps:").pack(anchor="w", pady=(0, 4))
        ctk.CTkCheckBox(right_panel, text="Copy output to 'Received by BBI'", variable=self.copy_output_var).pack(anchor="w", pady=2)
        ctk.CTkCheckBox(right_panel, text="Delete intermediate files when finished", variable=self.delete_intermediate_var).pack(anchor="w", pady=2)
        ctk.CTkCheckBox(right_panel, text="Replace plasmid ID with tube ID", variable=self.rename_from_excel_var).pack(anchor="w", pady=2)

    def validate_inputs(self) -> bool:
        """
        Validates that all inputs required for this specific workflow are ready.
        """
        try:
            float(self.max_w.get())
            float(self.max_h.get())
        except ValueError:
            messagebox.showerror("Error", "Max width/height must be numbers.")
            return False

        if self.copy_output_var.get():
            default_output_path = self.app.make_path_from_config('Paths', 'QC_Output_Folder')
            if not default_output_path or not os.path.isdir(default_output_path):
                messagebox.showerror("Error", "The destination folder for copying output files ('Received by BBI') does not exist.")
                return False
        
        if self.rename_from_excel_var.get():
            shipment_qc_path = self.app.make_path_from_config('Paths', 'Shipment_and_QC_Log')
            if not shipment_qc_path or not os.path.isfile(shipment_qc_path):
                messagebox.showerror("Error", "The 'Shipment and QC.xlsx' file was not found. Please check config.ini or uncheck the rename option.")
                return False
        return True

    def run_processing_task(self):
        """
        The main processing logic for the InHouse NewNames workflow.
        """
        self.log("[INFO] \t Starting InHouse NewNames workflow...")
        input_folder = os.path.normpath(self.app.input_folder.get().strip())

        if self.rename_from_excel_var.get():
            self.load_shipment_qc_data()

        extracted_zips_dir, extracted_zips_names = self.extract_zip_files(input_folder)
        if not extracted_zips_dir:
            self.log("[FINISHED] \t No files to process.")
            self.app.process_button.configure(state="normal")
            return

        stats = {
            'X': {'success': 0, 'fail': 0},
            'Z': {'success': 0, 'fail': 0},
            'SEQ': {'success_gel': 0, 'success_fallback': 0, 'fail': 0},
            'unrecognized': 0
        }
        total_files = len(extracted_zips_names)
        processed_folders_info = [] # Will store dicts: {'name': str, 'job_id': str, 'plasmid_id': str}

        for folder_name in extracted_zips_names:
            folder_dir = os.path.join(extracted_zips_dir, folder_name)

            mode = None
            if folder_name.startswith(self.NO_GEL_PREFIX):
                mode = self.NO_GEL_PREFIX # No gel processing
            elif folder_name.startswith(self.GEL_PREFIX):
                mode = self.GEL_PREFIX
            elif folder_name.startswith(self.SEQ_PREFIX):
                mode = self.SEQ_PREFIX

            if mode:
                status = self.process_single_folder(folder_dir, mode=mode)
                is_success = status['status'] in [ProcessStatus.SUCCESS_GEL, ProcessStatus.SUCCESS_NO_GEL, ProcessStatus.SUCCESS_FALLBACK_NO_GEL]

                if is_success:
                    processed_folders_info.append({'name': folder_name, 'job_id': status.get('job_id'), 'plasmid_id': status.get('plasmid_id')})

                if mode == self.NO_GEL_PREFIX:
                    if is_success: stats['X']['success'] += 1
                    else: stats['X']['fail'] += 1
                elif mode == self.GEL_PREFIX:
                    if is_success: stats['Z']['success'] += 1
                    else: stats['Z']['fail'] += 1
                elif mode == self.SEQ_PREFIX:
                    if status['status'] == ProcessStatus.SUCCESS_GEL: stats['SEQ']['success_gel'] += 1
                    elif status['status'] == ProcessStatus.SUCCESS_FALLBACK_NO_GEL: stats['SEQ']['success_fallback'] += 1
                    else: stats['SEQ']['fail'] += 1
            else:
                stats['unrecognized'] += 1
                self.log(f"[SKIP] \t Folder '{folder_name}' does not start with a recognized prefix (X, Z, SEQ).")
        
        # --- Rename folders based on Excel data before zipping ---
        processed_folder_names = [info['name'] for info in processed_folders_info]
        if self.rename_from_excel_var.get() and self.shipment_qc_df is not None:
            processed_folder_names = self.rename_processed_folders(extracted_zips_dir, processed_folders_info)

        zip_output_dir = os.path.join(os.path.dirname(input_folder), self.OUTPUT_ZIPS_DIR_NAME)
        if processed_folder_names:
            # Pass the (potentially renamed) list of folder names to be zipped
            self.zip_processed_folders(extracted_zips_dir, zip_output_dir, processed_folder_names)

        # Copy and Delete steps should use the final (renamed) folder names
        if self.copy_output_var.get() and processed_folder_names:
            default_output_path = self.app.make_path_from_config('Paths', 'QC_Output_Folder')
            self.log("[OK] \t Copying output files to default location...")
            for zip_name in processed_folder_names:
                source_zip_path = os.path.join(zip_output_dir, f"{zip_name}.zip")
                if os.path.exists(source_zip_path):
                    try:
                        shutil.copy(source_zip_path, default_output_path)
                        self.log(f"[OK] \t Copied {zip_name}.zip to {default_output_path}")
                    except Exception as e:
                        self.log(f"[ERROR] \t Failed to copy {zip_name}.zip: {e}")

        if self.delete_intermediate_var.get() and processed_folder_names:
            self.log("[OK] \t Deleting intermediate files as requested...")
            for dir_path in [extracted_zips_dir, zip_output_dir]:
                if os.path.isdir(dir_path):
                    try:
                        shutil.rmtree(dir_path)
                        self.log(f"[OK] \t Deleted intermediate folder: {dir_path}")
                    except Exception as e:
                        self.log(f"[ERROR] \t Could not delete {dir_path}: {e}")

        # --- Final Summary ---
        total_processed = len(processed_folder_names)
        summary_header = f"Processed {total_processed}/{total_files} files"
        summary_details = []

        x_success, x_fail = stats['X']['success'], stats['X']['fail']
        if x_success > 0 or x_fail > 0:
            summary_details.append(f"{x_success} 'X' files ({x_fail} failed)")

        z_success, z_fail = stats['Z']['success'], stats['Z']['fail']
        if z_success > 0 or z_fail > 0:
            summary_details.append(f"{z_success} 'Z' files ({z_fail} skipped/failed)")

        seq_gel, seq_fallback, seq_fail = stats['SEQ']['success_gel'], stats['SEQ']['success_fallback'], stats['SEQ']['fail']
        seq_success = seq_gel + seq_fallback
        if seq_success > 0 or seq_fail > 0:
            detail = f"{seq_success} 'SEQ' files"
            sub_details = [s for s in [f"{seq_gel} w/ gel" if seq_gel > 0 else "",
                                        f"{seq_fallback} no gel" if seq_fallback > 0 else "",
                                        f"{seq_fail} failed" if seq_fail > 0 else ""] if s]
            if sub_details:
                detail += f" ({', '.join(sub_details)})"
            summary_details.append(detail)

        unrec = stats['unrecognized']
        if unrec > 0:
            summary_details.append(f"{unrec} unrecognized")

        if summary_details:
            final_summary = f"{summary_header}: {'; '.join(summary_details)}."
        else:
            final_summary = summary_header + "."

        self.log(f"[SUMMARY] \t {final_summary}")
        self.log("[FINISHED] \t InHouse NewNames workflow complete!")
        self.app.process_button.configure(state="normal")

    # --- Helper Methods for this Workflow ---

    def load_shipment_qc_data(self):
        """Loads the 'Label_Output' sheet from the Shipment QC Excel file."""
        shipment_qc_path = self.app.make_path_from_config('Paths', 'Shipment_and_QC_Log')
        if not shipment_qc_path or not os.path.isfile(shipment_qc_path):
            self.log("[ERROR] \t Shipment QC Excel file not found. Cannot perform renaming.")
            self.shipment_qc_df = None
            return

        try:
            self.log("[INFO] \t Loading Shipment QC data for renaming...")
            self.shipment_qc_df = pd.read_excel(shipment_qc_path, sheet_name="Label_Output", header=0)
            # Ensure required columns are strings for reliable matching
            self.shipment_qc_df['Tube ID'] = self.shipment_qc_df['Tube ID'].astype(str)
            self.shipment_qc_df['BBI Lot#'] = self.shipment_qc_df['BBI Lot#'].astype(str)
            self.log("[OK] \t Successfully loaded Shipment QC data.")
        except Exception as e:
            self.log(f"[ERROR] \t Failed to load or read 'Shipment and QC.xlsx': {e}")
            self.shipment_qc_df = None

    def rename_processed_folders(self, base_dir: str, processed_folders: List[dict]) -> List[str]:
        """Renames folders based on data from the loaded Excel sheet."""
        renamed_folders = []
        for folder_info in processed_folders:
            original_name = folder_info['name']
            plasmid_id = folder_info['plasmid_id']
            
            if not plasmid_id:
                self.log(f"[WARNING] \t Could not find Plasmid ID for '{original_name}'. Skipping rename.")
                renamed_folders.append(original_name)
                continue

            try:
                parts = original_name.split(' ')
                match = self.shipment_qc_df[
                    (self.shipment_qc_df['BBI Lot#'] == plasmid_id)
                ]

                if not match.empty:
                    strain = str(match['Strain'].iloc[0])
                    strain_suffix = strain[1:] if len(strain) > 1 else strain
                    base_name_without_plasmid = original_name.rsplit('.', 1)[0]
                    new_name = f"{base_name_without_plasmid} {strain_suffix}"
                    
                    # Clean up potential old plasmid ID formats from the original name
                    base_name_without_plasmid = original_name.replace(f".{plasmid_id}", "").replace(f" {plasmid_id}", "").strip()
                    new_name = f"{base_name_without_plasmid} {strain_suffix}"
                    
                    original_outer_path = os.path.join(base_dir, original_name)
                    new_outer_path = os.path.join(base_dir, new_name)
                    original_inner_path = os.path.join(original_outer_path, original_name)
                    new_inner_path = os.path.join(original_outer_path, new_name)

                    os.rename(original_inner_path, new_inner_path) # Rename inner folder first
                    os.rename(original_outer_path, new_outer_path) # Then rename outer folder
                    renamed_folders.append(new_name)
                    self.log(f"[OK] \t Renamed '{original_name}' to '{new_name}'")
                else:
                    self.log(f"[WARNING] \t No match found in Excel for Plasmid ID '{plasmid_id}'. Removing plasmid ID from name.")
                    # If no match, just remove the plasmid ID part from the name
                    new_name = original_name.replace(f".{plasmid_id}", "").replace(f" {plasmid_id}", "").strip()

                    original_outer_path = os.path.join(base_dir, original_name)
                    new_outer_path = os.path.join(base_dir, new_name)
                    original_inner_path = os.path.join(original_outer_path, original_name)
                    new_inner_path = os.path.join(original_outer_path, new_name)

                    os.rename(original_inner_path, new_inner_path)
                    os.rename(original_outer_path, new_outer_path)
                    renamed_folders.append(new_name)
            except Exception as e:
                self.log(f"[ERROR] \t Failed to rename folder '{original_name}': {e}")
                renamed_folders.append(original_name)
        return renamed_folders

    def setup_default_paths(self):
        default_path = self.app.make_path_from_config('Paths', 'QC_Input_Folder')
        if default_path and os.path.isdir(default_path):
            self.app.input_folder.set(default_path)
            self.log(f"[OK] \t Default input folder found and set to: {default_path}")

    def extract_zip_files(self, input_folder: str) -> Tuple[str, List[str]]:
        extracted_folders_dir = os.path.join(os.path.dirname(input_folder), self.PROCESSED_FOLDERS_DIR_NAME)
        if os.path.exists(extracted_folders_dir):
            shutil.rmtree(extracted_folders_dir)
        os.makedirs(extracted_folders_dir)
        self.log(f"[OK] \t Created clean processing folder: {extracted_folders_dir}")

        zip_files = [f for f in os.listdir(input_folder) if f.lower().endswith('.zip')]
        if not zip_files:
            return None, []

        for zip_file in zip_files:
            zip_path = os.path.join(input_folder, zip_file)
            # This workflow expects a nested folder structure, so we extract into a sub-folder
            extract_to = os.path.join(extracted_folders_dir, os.path.splitext(zip_file)[0])
            try:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(extract_to)
                    self.log(f"[OK] \t Extracted {zip_file}")
            except Exception as e:
                self.log(f"[ERROR] \t Failed to extract {zip_file}: {e}")

        return extracted_folders_dir, [os.path.splitext(f)[0] for f in zip_files]

    def find_report_document(self, directory: str) -> str:
        """
        Finds the most likely report document (.doc or .docx) in a given directory.
    
        Logic:
        1. Find all .doc and .docx files, ignoring temporary files.
        2. If exactly one is found, return its path.
        3. If multiple are found, filter by report-specific substrings.
        4. If filtering results in one match, return its path.
        5. If still multiple (or none from filter), sort alphabetically and return the first one.
        6. If no Word documents are found initially, return None.
        """
        all_word_docs = [f for f in os.listdir(directory)
                        if (f.lower().endswith(self.DOC_EXTENSION) or f.lower().endswith(self.DOCX_EXTENSION))
                        and not f.startswith(self.INVALID_FILE_PREFIX)]
    
        if not all_word_docs:
            self.log(f"[ERROR] \t No Word documents (.doc or .docx) found in {directory}")
            return None
    
        if len(all_word_docs) == 1:
            self.log(f"[OK] \t Found single Word document: {all_word_docs[0]}")
            return os.path.join(directory, all_word_docs[0])
    
        # If multiple documents, try to find the best match
        self.log(f"[INFO] \t Multiple Word documents found. Attempting to select the correct report.")
        
        spaceless_substring = self.REPORT_SUBSTRING.replace(' ', '')
        candidates = [f for f in all_word_docs 
                        if self.REPORT_SUBSTRING in f.lower() or spaceless_substring in f.lower()]
    
        if len(candidates) == 1:
            self.log(f"[OK] \t Selected report by substring match: {candidates[0]}")
            return os.path.join(directory, candidates[0])
        
        final_list = candidates if candidates else all_word_docs
        final_list.sort()
        chosen_doc = final_list[0]
        self.log(f"[WARNING] \t Multiple candidates or no substring match. Selecting first alphabetically: {chosen_doc}")
        return os.path.join(directory, chosen_doc)

    def process_single_folder(self, folder_dir: str, mode: str) -> dict:
        base_name = os.path.basename(folder_dir)
        docx_subfolder_dir = os.path.join(folder_dir, base_name)
        if not os.path.isdir(docx_subfolder_dir):
            self.log(f"[ERROR] \t Expected subfolder not found: {docx_subfolder_dir}")
            return {'status': ProcessStatus.FAIL_SUBFOLDER_NOT_FOUND}

        doc_path = self.find_report_document(docx_subfolder_dir)
        if not doc_path:
            return {'status': ProcessStatus.FAIL_DOC_NOT_FOUND}
        
        is_doc_format = doc_path.lower().endswith(self.DOC_EXTENSION)
        word_app = None
        temp_docx_path = None
        com_initialized = False

        try:
            # Initialize COM for this thread (needed for .doc conversion OR .docx header fix)
            pythoncom.CoInitialize()
            com_initialized = True

            doc_obj = None
            if is_doc_format:
                # Use Dispatch instead of EnsureDispatch for better compatibility across Office versions.
                word_app = win32.Dispatch('Word.Application')
                word_app.Visible = False
                temp_docx_path = os.path.splitext(doc_path)[0] + self.TEMP_DOCX_SUFFIX
                
                doc = word_app.Documents.Open(doc_path)
                # Use SaveAs for older Word versions that may not have SaveAs2
                if hasattr(doc, 'SaveAs2'):
                    doc.SaveAs2(temp_docx_path, FileFormat=self.WIN32_DOCX_FORMAT)
                else:
                    doc.SaveAs(temp_docx_path, FileFormat=self.WIN32_DOCX_FORMAT)
                doc.Close(False)
                self.log(f"[OK] \t Converted .doc to temporary .docx for editing.")
                doc_obj = Document(temp_docx_path)
            else:
                doc_obj = Document(doc_path)

            for section in doc_obj.sections:
                section.footer_distance = Cm(1.0)

            # --- NEW: Extract IDs from Word doc first ---
            job_id, plasmid_id = self.extract_ids_from_doc(doc_obj)
            if not job_id or not plasmid_id:
                self.log(f"[ERROR] \t Could not extract Job ID or Plasmid ID from {os.path.basename(doc_path)}")
                return {'status': ProcessStatus.FAIL_IDS_NOT_FOUND}
            self.log(f"[INFO] \t Extracted from doc: JobID='{job_id}', PlasmidID='{plasmid_id}'")

            status = None
            if mode == self.NO_GEL_PREFIX:
                self.remove_qc_elements(doc_obj, base_name)
                status = ProcessStatus.SUCCESS_NO_GEL

            elif mode == self.GEL_PREFIX:
                image_found = self.find_and_insert_gel_image(job_id, plasmid_id, doc_obj)
                if not image_found:
                    return {'status': ProcessStatus.FAIL_NO_GEL}  # Fail hard for 'Z' files
                status = ProcessStatus.SUCCESS_GEL

            elif mode == self.SEQ_PREFIX:
                image_found = self.find_and_insert_gel_image(job_id, plasmid_id, doc_obj)
                if not image_found:
                    self.log(f"[INFO] \t No gel for SEQ file '{base_name}'. Removing QC table as fallback.")
                    self.remove_qc_elements(doc_obj, base_name)
                    status = ProcessStatus.SUCCESS_FALLBACK_NO_GEL
                else:
                    status = ProcessStatus.SUCCESS_GEL
            
            if is_doc_format:
                doc_obj.save(temp_docx_path)
                doc = word_app.Documents.Open(temp_docx_path)
                if hasattr(doc, 'SaveAs2'):
                    doc.SaveAs2(doc_path, FileFormat=self.WIN32_DOC_FORMAT)
                else:
                    doc.SaveAs(doc_path, FileFormat=self.WIN32_DOC_FORMAT)
                doc.Close(False)
                self.log(f"[OK] \t Saved changes back to original .doc format.")
            else:
                doc_obj.save(doc_path)
                # Apply header fix for .docx files as requested
                self.fix_header_formatting(doc_path)

            return {'status': status, 'job_id': job_id, 'plasmid_id': plasmid_id}
        except Exception as e:
            self.log(f"[ERROR] \t Failed to process Word document {doc_path}: {e}")
            return {'status': ProcessStatus.FAIL_UNEXPECTED}
        finally:
            if temp_docx_path and os.path.exists(temp_docx_path):
                try:
                    os.remove(temp_docx_path)
                except OSError as e:
                    self.log(f"[WARNING] \t Could not remove temp file {temp_docx_path}: {e}")
            if word_app:
                word_app.Quit()
            if com_initialized:
                # Uninitialize COM for this thread
                pythoncom.CoUninitialize()

    def fix_header_formatting(self, doc_path: str):
        """
        Fixes header formatting for .docx files using Word COM:
        - Identifies the widest shape as the background rectangle.
        - Rectangle: Shift right 20, Behind Text, Send to Back.
        - Others: In Front of Text, Bring to Front.
        - Sets entire document font to Calibri.
        """
        word_app = None
        temp_path = None
        try:
            # Copy to a temporary file to bypass Word's path length limitations
            fd, temp_path = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            shutil.copy2(doc_path, temp_path)

            word_app = win32.Dispatch('Word.Application')
            word_app.Visible = False
            doc = word_app.Documents.Open(temp_path)
            
            modified = False
            for section in doc.Sections:
                for header in section.Headers:
                    if header.Exists:
                        # Collect all shapes in the header
                        shapes = []
                        for shape in header.Shapes:
                            shapes.append(shape)
                        
                        if shapes:
                            # Sort by width descending. The widest is likely the background rectangle.
                            try:
                                shapes.sort(key=lambda s: s.Width, reverse=True)
                                
                                # Background Rectangle (Widest)
                                bg_shape = shapes[0]
                                bg_shape.Left += 2
                                bg_shape.WrapFormat.Type = 5 # wdWrapBehind
                                bg_shape.ZOrder(1) # msoSendToBack
                                
                                # Foreground Objects (Others)
                                for shape in shapes[1:]:
                                    shape.WrapFormat.Type = 3 # wdWrapFront
                                    shape.ZOrder(0) # msoBringToFront
                                
                                modified = True
                            except Exception:
                                pass
            
            # Homogenize font to Calibri (Ctrl+A equivalent)
            doc.Content.Font.Name = "Calibri"
            modified = True

            if modified:
                doc.Save()
            doc.Close(False)

            if modified:
                shutil.copy2(temp_path, doc_path)
                self.log(f"[OK] \t Applied header layout fix to {os.path.basename(doc_path)}")
        except Exception as e:
            self.log(f"[WARNING] \t Failed to apply header fix: {e}")
        finally:
            if word_app:
                try:
                    word_app.Quit()
                except:
                    pass
            if temp_path and os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass

    def extract_ids_from_doc(self, doc: Document) -> Tuple[str, str]:
        """Extracts Job ID and Plasmid ID from the first table of the document."""
        job_id, plasmid_id = None, None
        if not doc.tables:
            return None, None
        
        table = doc.tables[0]
        for row in table.rows:
            # Check if row has at least 2 cells to avoid IndexError
            if len(row.cells) < 2:
                continue
            
            label = row.cells[0].text.strip().lower()
            if "work no." in label:
                job_id = row.cells[1].text.strip()
            elif "plasmid no." in label:
                plasmid_id = row.cells[1].text.strip()
        return job_id, plasmid_id

    def remove_qc_elements(self, doc: Document, base_name: str):
        """Removes the QC table and its preceding paragraph from the document."""
        table_removed = False
        if len(doc.tables) > self.QC_TABLE_INDEX:
            tbl = doc.tables[self.QC_TABLE_INDEX]
            tbl._element.getparent().remove(tbl._element)
            table_removed = True

        para_removed = False
        para_to_delete = next((p for p in doc.paragraphs if self.QC_TABLE_PARAGRAPH_TEXT in p.text.lower()), None)
        if para_to_delete:
            p_element = para_to_delete._element
            p_element.getparent().remove(p_element)
            para_removed = True
        
        if table_removed or para_removed:
            self.log(f"[OK] \t Removed QC elements from docx for {base_name}")
        else:
            self.log(f"[INFO] \t No QC elements found to remove in {base_name}. Document may be pre-processed.")

    def get_year_letter(self, year: int) -> str:
        """
        Generates the year-code letter based on the year (e.g., S=2024, T=2025).
        This is based on the folder naming convention described.
        """
        base_year = 2024
        base_letter_ord = ord('S')
        if year < base_year:
            self.log(f"[WARNING] \t Cannot generate year letter for year {year} (before {base_year}).")
            return None
        return chr(base_letter_ord + (year - base_year))

    def get_recent_folder_info(self, months_to_check: int = 6) -> Tuple[set, set]:
        """
        Generates sets of valid folder prefixes and relevant years for recent months.
        This limits the search space for gel images to improve performance.
        e.g., for a run in Feb 2026, returns:
        ({'U02', 'U01', 'T12', 'T11', 'T10', 'T09'}, {'2026', '2025'})
        """
        today = date.today()
        valid_prefixes = set()
        valid_years = set()

        current_year = today.year
        current_month = today.month

        for _ in range(months_to_check):
            year_letter = self.get_year_letter(current_year)
            if year_letter:
                # Subfolders are named LMMDD, so we create a prefix LMM to check against.
                prefix = f"{year_letter}{current_month:02d}"
                valid_prefixes.add(prefix)
                valid_years.add(str(current_year))

            # Move to the previous month
            current_month -= 1
            if current_month == 0:
                current_month = 12
                current_year -= 1
        
        self.log(f"[INFO] \t Limiting gel image search to recent folders starting with: {sorted(list(valid_prefixes), reverse=True)}")
        return valid_prefixes, valid_years

    def find_and_insert_gel_image(self, job_id: str, plasmid_id: str, doc: Document) -> bool:
        gel_images_folder = self.app.make_path_from_config('Paths', 'QC_Gel_Images')
        if not gel_images_folder or not os.path.isdir(gel_images_folder):
            self.log(f"[ERROR] \t Gel images folder not found: {gel_images_folder}")
            return False

        try:
            # To improve performance, only search in folders from the last 6 months.
            recent_folder_prefixes, recent_years = self.get_recent_folder_info()

            all_year_folders = [f for f in os.listdir(gel_images_folder) if os.path.isdir(os.path.join(gel_images_folder, f))]
            year_folders_to_search = sorted([f for f in all_year_folders if f in recent_years], reverse=True)

        except FileNotFoundError:
            self.log(f"[ERROR] \t Gel images folder does not exist: {gel_images_folder}")
            return False

        for year_folder in year_folders_to_search:
            year_path = os.path.join(gel_images_folder, year_folder)
            # Filter subfolders to only include those matching the recent prefixes.
            subfolders = sorted([f for f in os.listdir(year_path)
                                if os.path.isdir(os.path.join(year_path, f))
                                and len(f) >= 5 and f[:3] in recent_folder_prefixes],
                                reverse=True)
            for subfolder in subfolders:
                subfolder_path = os.path.join(year_path, subfolder)
                for fname in os.listdir(subfolder_path):
                    if not fname.lower().endswith(('.jpg', '.jpeg', '.png')): continue
                    img_name_no_ext = os.path.splitext(fname)[0]
                    parts = img_name_no_ext.split('.', 3)
                    if len(parts) != 4: continue

                    job_id_from_img, plasmid_id_from_img, enzymes, expected_size = parts
                    # Use exact match now that we have the IDs from the Word doc
                    if job_id_from_img == job_id and plasmid_id_from_img == plasmid_id:
                        self.log(f"[OK] \t Found gel image: {fname}")
                        self.log(f"[OK] \t Match: JobID '{job_id_from_img}' & PlasmidID '{plasmid_id_from_img}'")
                        self.log(f"[OK] \t Enzyme: {enzymes} | Expected Size: {expected_size}")
                        
                        image_path = os.path.join(subfolder_path, fname)
                        self.insert_image_into_doc(doc, image_path, enzymes, expected_size)
                        return True

        self.log(f"[ERROR] \t No matching gel image found for JobID '{job_id}' and PlasmidID '{plasmid_id}' (searched last 6 months).")
        return False

    def insert_image_into_doc(self, doc: Document, image_path: str, enzymes: str, expected_size: str):
        if len(doc.tables) <= self.QC_TABLE_INDEX:
            self.log(f"[ERROR] \t Word document has fewer than {self.QC_TABLE_INDEX + 1} tables. Cannot insert image.")
            return
        table = doc.tables[self.QC_TABLE_INDEX]
        table.allow_autofit = False
        
        target_row_idx = next((idx for idx, row in enumerate(table.rows) if len(row.cells) >= 2 and ("enzyme" in row.cells[1].text.lower() or "expected size" in row.cells[1].text.lower())), len(table.rows) - 1)
        row = table.rows[target_row_idx]
        if len(row.cells) < 2:
            self.log(f"[ERROR] \t Target row has fewer than 2 cells.")
            return

        row.height = Cm(8.7)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # Insert gel image
        gel_cell = row.cells[0]
        gel_cell.text = ""
        gel_paragraph = gel_cell.paragraphs[0]
        gel_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            max_w_in, max_h_in = float(self.max_w.get()), float(self.max_h.get())
        except ValueError:
            max_w_in, max_h_in = float(self.DEFAULT_MAX_IMG_WIDTH), float(self.DEFAULT_MAX_IMG_HEIGHT)
        
        with Image.open(image_path) as im:
            width, height = im.size
        aspect = width / height
        
        width_in, height_in = (Inches(max_w_in), None) if (max_w_in / aspect) <= max_h_in else (None, Inches(max_h_in))
        
        run = gel_paragraph.add_run()
        run.add_picture(image_path, width=width_in, height=height_in)

        # Insert reference paragraph
        self.write_ref_paragraph(row.cells[1], enzymes, expected_size)
        self.log(f"[OK] \t Inserted gel image and info into document.")
        
    def write_ref_paragraph(self, info_cell, enzyme: str, expected_size: str):
        info_cell.text = ""
        p = info_cell.paragraphs[0]
        
        def add_formatted_line(label, value):
            run_label = p.add_run(label)
            run_label.bold = True
            run_label.font.color.rgb = self.BBI_BLUE_COLOR
            p.add_run().add_break()
            p.add_run(value)
            p.add_run().add_break()

        add_formatted_line("Enzyme:", enzyme)
        add_formatted_line("Expected Size:", expected_size)

        ladder_label = p.add_run("Marker Ladder:")
        ladder_label.bold = True
        ladder_label.font.color.rgb = self.BBI_BLUE_COLOR
        p.add_run().add_break()

        # Try path passed from hub first, then fall back to config file
        marker_img_path = self.app.hub_paths.get('Gel_Ladder')
        if not marker_img_path or not os.path.isfile(marker_img_path):
            marker_img_path = self.app.make_path_from_config('Paths', 'QC_Gel_Ladder')
            
        if marker_img_path:
            try:
                p.add_run().add_picture(marker_img_path, height=Inches(2.4))
            except Exception as e:
                self.log(f"[ERROR] \t Could not insert marker ladder image: {e}")
        else:
            self.log("[WARNING] \t No marker ladder image found in either configured path.")

    def zip_processed_folders(self, extracted_zips_dir: str, zip_output_dir: str, folder_names: List[str]):
        if not os.path.exists(zip_output_dir):
            os.makedirs(zip_output_dir)

        for folder_name in folder_names:
            # The content to be zipped is in a nested subfolder
            outer_folder_path = os.path.join(extracted_zips_dir, folder_name)
            inner_folder_to_zip = os.path.join(outer_folder_path, folder_name)
            output_zip_base_path = os.path.join(zip_output_dir, folder_name)
            
            if os.path.isdir(inner_folder_to_zip):
                try:
                    # We zip the contents of the outer folder, which is just the inner folder
                    shutil.make_archive(output_zip_base_path, 'zip', root_dir=outer_folder_path, base_dir=folder_name)
                    self.log(f"[OK] \t Zipped {folder_name} to {output_zip_base_path}.zip")
                except Exception as e:
                    self.log(f"[ERROR] \t Failed to zip {folder_name}: {e}")
            else:
                self.log(f"[ERROR] \t Inner processed folder not found for zipping: {inner_folder_to_zip}")
