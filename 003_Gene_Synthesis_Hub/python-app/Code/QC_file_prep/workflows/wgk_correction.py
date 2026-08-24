"""
Bio Basic Inc. - Canada
Made by: Eduardo Reyes, Ph.D.
Contact: ed5reyes@outlook

Version: 1.1
Date: Mar 30, 2026
Notes: Fixed issue with hub-passed directories.
"""

import os
import shutil
import sys
import zipfile
from typing import List, Optional, Tuple

import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.table import Table
import pythoncom
import win32com.client as win32

class WgkCorrection:
    def __init__(self, app_instance):
        """
        Initialize the WGK Correction workflow.
        - app_instance: A reference to the main App to access common elements.
        """
        self.app = app_instance
        # Use the main application's logging instance for consistent logging
        self.log = self.app.log

        # --- Constants for this workflow ---

        # Directory names for intermediate and final output
        self.PROCESSED_FOLDERS_DIR_NAME = "2-processed_folders_wgk"
        self.OUTPUT_ZIPS_DIR_NAME = "3-output_zips_wgk"

        # Suffix for the default input folder specific to this workflow
        self.DEFAULT_PATH_SUFFIX = "_wgk"

        # UI modes for input
        self.SINGLE_MODE = "single"
        self.BATCH_MODE = "batch"

        # Default value and file format for batch processing
        self.DEFAULT_PLASMID_QTY = "10 µg"
        self.BATCH_FILE_DELIMITER = '\t'

        # File handling constants
        self.DOC_EXTENSION = ".doc"
        self.DOCX_EXTENSION = ".docx"
        self.INVALID_FILE_PREFIX = '~'
        self.REPORT_SUBSTRING = "gene synthesis report"
        self.TEMP_DOCX_SUFFIX = "_temp.docx"

        # Constants for COM interaction with MS Word
        self.WIN32_DOC_FORMAT = 0  # wdFormatDocument
        self.WIN32_DOCX_FORMAT = 16 # wdFormatXMLDocument

        # Constants for locating and editing the target cell in the Word document
        self.TARGET_CELL_TEXT = "plasmid qty"
        self.TARGET_CELL_INDEX = 0
        self.QTY_CELL_INDEX = 2
        self.QTY_PADDING = "   "

        # Default font settings for the new text
        self.DEFAULT_FONT_NAME = 'Arial'
        self.DEFAULT_FONT_SIZE = 9

        # --- Workflow-specific variables ---

        # CTkinter variables to hold the state of the UI widgets
        self.input_mode = ctk.StringVar(value=self.SINGLE_MODE)
        self.plasmid_qty = ctk.StringVar(value=self.DEFAULT_PLASMID_QTY)
        self.batch_file_path = ctk.StringVar()
        self.copy_output_var = ctk.BooleanVar(value=False)
        self.delete_intermediate_var = ctk.BooleanVar(value=False)

        # Dictionary to store mappings from the batch file (ID -> new value)
        self.batch_replacements = {}

    def create_ui(self, parent_frame):
        """
        Creates the UI elements specific to this workflow inside the provided parent frame.
        """
        parent_frame.grid_columnconfigure(0, weight=1)

        # --- Main panel for UI elements ---
        main_panel = ctk.CTkFrame(parent_frame, fg_color="transparent")
        main_panel.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)

        # Configure a 65/35 proportional layout for the main panel
        main_panel.grid_columnconfigure(0, weight=13) # 65%
        main_panel.grid_columnconfigure(1, weight=7)  # 35%

        # --- Left Column for Inputs ---
        left_panel = ctk.CTkFrame(main_panel, fg_color="transparent")
        left_panel.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        # --- Right Column for Options ---
        right_panel = ctk.CTkFrame(main_panel, fg_color="transparent")
        right_panel.grid(row=0, column=1, sticky="nsew", padx=(10, 0))

        # --- Populate Left Panel ---

        # Input Mode Selection
        mode_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        mode_frame.pack(pady=8, anchor="w")
        ctk.CTkLabel(mode_frame, text="Input Mode:").pack(side="left", padx=(0, 15))
        ctk.CTkRadioButton(mode_frame, text="Single Value", variable=self.input_mode, value=self.SINGLE_MODE, command=self.toggle_input_mode).pack(side="left", padx=5)
        ctk.CTkRadioButton(mode_frame, text="Batch from File", variable=self.input_mode, value=self.BATCH_MODE, command=self.toggle_input_mode).pack(side="left", padx=5)

        # Single Value Input Frame
        self.single_mode_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        ctk.CTkLabel(self.single_mode_frame, text="New Plasmid Qty Value:").pack(side="left", padx=(0, 8))
        self.qty_entry = ctk.CTkEntry(self.single_mode_frame, textvariable=self.plasmid_qty, width=380)
        self.qty_entry.pack(side="left", padx=(0, 8))

        # Batch File Input Frame
        self.batch_mode_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        ctk.CTkLabel(self.batch_mode_frame, text="Batch File (.txt):").pack(side="left", padx=(0, 8))
        self.batch_file_entry = ctk.CTkEntry(self.batch_mode_frame, textvariable=self.batch_file_path, width=380)
        self.batch_file_entry.pack(side="left", padx=(0, 8))
        ctk.CTkButton(self.batch_mode_frame, text="Browse...", width=80, command=self.browse_batch_file).pack(side="left")

        # --- Populate Right Panel ---

        ctk.CTkLabel(right_panel, text="Optional Steps:").pack(anchor="w", pady=(0, 4))
        ctk.CTkCheckBox(right_panel, text="Copy output to 'Received by BBI'", variable=self.copy_output_var).pack(anchor="w", pady=2)
        ctk.CTkCheckBox(right_panel, text="Delete intermediate files when finished", variable=self.delete_intermediate_var).pack(anchor="w", pady=2)

        # --- Set initial UI state ---
        # Ensure the correct input frame is visible when the UI is first created
        self.toggle_input_mode()

    def validate_inputs(self) -> bool:
        """
        Validates that all inputs required for this specific workflow are ready.
        """
        mode = self.input_mode.get()

        # Validate inputs for "Single Value" mode
        if mode == self.SINGLE_MODE:
            if not self.plasmid_qty.get().strip():
                messagebox.showerror("Input Error", "The 'New Plasmid Qty Value' cannot be empty.")
                return False
        # Validate inputs for "Batch from File" mode
        elif mode == self.BATCH_MODE: 
            batch_file = self.batch_file_path.get()
            if not batch_file:
                messagebox.showerror("Input Error", "Please select a batch file.")
                return False
            if not os.path.exists(batch_file):
                messagebox.showerror("File Not Found", f"The specified batch file does not exist:\n{batch_file}")
                return False
            if not self.load_batch_replacements():
                # Error message is shown within the loading function
                return False

        # This workflow uses win32com to handle .doc files, which is Windows-only
        if sys.platform != "win32":
            messagebox.showerror("Unsupported OS",
                                "This workflow requires Windows to process .doc files.")
            return False

        # If the user wants to copy files, check if the destination path is valid
        if self.copy_output_var.get():
            # The path is retrieved from the shared config file
            default_output_path = self.app.make_path_from_config('Paths', 'QC_Output_Folder')
            if not default_output_path or not os.path.isdir(default_output_path):
                messagebox.showerror("Error",
                                    "The destination folder ('Received by BBI') does not exist. "
                                    "Please check default_directories.ini or uncheck the copy option.")
                return False
        return True

    def run_processing_task(self):
        """
        The main processing logic for the WGK Correction workflow.
        """
        self.log("[INFO] \t Starting WGK Correction workflow...")
        input_folder = os.path.normpath(self.app.input_folder.get().strip())

        # Get a list of all zip files in the selected input directory
        all_zip_files = [f for f in os.listdir(input_folder) if f.lower().endswith('.zip')]

        if self.input_mode.get() == self.BATCH_MODE:
            # In batch mode, only process zip files that have a corresponding entry in the batch file
            relevant_zips = []
            for zip_file in all_zip_files:
                zip_id = os.path.splitext(zip_file)[0]
                if any(zip_id.startswith(key) for key in self.batch_replacements.keys()):
                    relevant_zips.append(zip_file)
            
            if not relevant_zips:
                self.log("[WARNING] \t No zip files found matching the identifiers in the batch file.")
            all_zip_files = relevant_zips

        # Extract all relevant zip files to a temporary processing directory
        extracted_zips_dir, extracted_zips_names = self.extract_zip_files(input_folder, all_zip_files)
        if not extracted_zips_dir:
            self.log("[FINISHED] \t No zip files to process.")
            self.app.process_button.configure(state="normal")
            return
        
        # Initialize counters for the summary log
        processed_count = 0
        skipped_count = 0

        for folder_name in extracted_zips_names:
            # Determine the correct replacement value based on the mode (single vs. batch)
            new_value = self.get_value_for_folder(folder_name)
            if new_value is None:
                self.log(f"[WARNING] \t No replacement value found for '{folder_name}'. Skipping.")
                skipped_count += 1
                continue
            
            current_folder_path = os.path.join(extracted_zips_dir, folder_name)
            # Find the specific Word document to modify within the extracted folder
            doc_path = self.find_report_document(current_folder_path)
            if doc_path:
                # Perform the main modification logic on the found document
                self.process_word_document(doc_path, new_value)
                processed_count += 1

        # After processing, re-zip the modified folders into an output directory
        zip_output_dir = os.path.join(os.path.dirname(input_folder), self.OUTPUT_ZIPS_DIR_NAME)
        if extracted_zips_names:
            self.zip_processed_folders(extracted_zips_dir, zip_output_dir, extracted_zips_names)

        # If the user selected the option, copy the final zip files to a shared/default location
        if self.copy_output_var.get():
            default_output_path = self.app.make_path_from_config('Paths', 'Default_Output_QC_Folder')
            self.log("[OK] \t Copying output files to default location...")
            for zip_name in extracted_zips_names:
                source_zip_path = os.path.join(zip_output_dir, f"{zip_name}.zip")
                if os.path.exists(source_zip_path):
                    try:
                        shutil.copy(source_zip_path, default_output_path)
                        self.log(f"[OK] \t Copied {zip_name}.zip to {default_output_path}")
                    except Exception as e:
                        self.log(f"[ERROR] \t Failed to copy {zip_name}.zip: {e}")

        # If the user selected the option, clean up by deleting the intermediate folders
        if self.delete_intermediate_var.get():
            self.log("[OK] \t Deleting intermediate files as requested...")
            for dir_path in [extracted_zips_dir, zip_output_dir]:
                if os.path.isdir(dir_path):
                    try:
                        shutil.rmtree(dir_path)
                        self.log(f"[OK] \t Deleted intermediate folder: {dir_path}")
                    except Exception as e:
                        self.log(f"[ERROR] \t Could not delete {dir_path}: {e}")

        # Log a final summary of the operation
        total_files = len(extracted_zips_names)
        self.log(f"[SUMMARY] \t Processed {processed_count}/{total_files} files ({skipped_count} skipped).")
        self.log("[FINISHED] \t WGK Correction process complete!")
        self.app.process_button.configure(state="normal")

    # --- Helper Methods for this Workflow ---

    def toggle_input_mode(self):
        """Shows and hides UI elements based on the selected input mode."""
        # This ensures only the relevant input fields are visible to the user
        if self.input_mode.get() == self.SINGLE_MODE:
            self.batch_mode_frame.pack_forget()
            self.single_mode_frame.pack(pady=8, anchor="w", after=self.single_mode_frame.master.winfo_children()[0])
        else:
            self.single_mode_frame.pack_forget()
            self.batch_mode_frame.pack(pady=8, anchor="w", after=self.batch_mode_frame.master.winfo_children()[0])

    def browse_batch_file(self):
        """Opens a file dialog to select a .txt batch file."""
        filepath = filedialog.askopenfilename(
            title="Select Batch File",
            filetypes=(("Text files", "*.txt"), ("All files", "*.*"))
        )
        if filepath:
            self.batch_file_path.set(filepath)

    def load_batch_replacements(self) -> bool:
        """Parses the batch file and loads data into a dictionary."""
        self.batch_replacements.clear()

        try:
            # Use 'utf-8' encoding to correctly read special characters like 'µ'
            with open(self.batch_file_path.get(), 'r', encoding='utf-8') as f:
                for i, line in enumerate(f, 1):
                    # Skip empty lines or lines without the delimiter
                    if not line.strip() or self.BATCH_FILE_DELIMITER not in line:
                        continue
                    # Split the line into key (file identifier) and value (new plasmid qty)
                    parts = line.strip().split(self.BATCH_FILE_DELIMITER, 1)
                    # Store in the dictionary for later lookup
                    self.batch_replacements[parts[0].strip()] = parts[1].strip()
            self.log(f"[OK] \t Loaded {len(self.batch_replacements)} replacements from batch file.")
            return True
        except Exception as e:
            messagebox.showerror("Batch File Error", f"Failed to read or parse the batch file.\nEnsure it is a tab-delimited UTF-8 text file.\n\nError: {e}")
            self.log(f"[ERROR] \t Failed to parse batch file: {e}")
            return False

    def setup_default_paths(self):
        """
        Sets up the default input folder for this workflow, creating it if it doesn't exist.
        """
        default_path = self.app.make_path_from_config('Paths', 'QC_Input_Folder')
        if default_path:
            self.app.input_folder.set(default_path)
            self.log(f"[OK] \t Default input folder set to: {default_path}")
            
    def extract_zip_files(self, input_folder: str, zip_files: List[str]) -> Tuple[str, List[str]]:
        """
        Extracts a list of zip files into a clean temporary directory.
        """
        extracted_dir = os.path.join(os.path.dirname(input_folder), self.PROCESSED_FOLDERS_DIR_NAME)

        # Ensure the extraction directory is clean before starting
        if os.path.exists(extracted_dir):
            shutil.rmtree(extracted_dir)
        os.makedirs(extracted_dir)
        self.log(f"[OK] \t Created clean processing folder: {extracted_dir}")

        if not zip_files:
            return None, []

        # Extract each zip file into the target directory
        for zip_file in zip_files:
            zip_path = os.path.join(input_folder, zip_file)
            try:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(extracted_dir)
                    self.log(f"[OK] \t Extracted {zip_file}")
            except Exception as e:
                self.log(f"[ERROR] \t Failed to extract {zip_file}: {e}")

        # Return the path to the extraction directory and a list of the folder names (from zip names)
        return extracted_dir, [os.path.splitext(f)[0] for f in zip_files]

    def get_value_for_folder(self, folder_name: str) -> str:
        """Gets the correct replacement value based on the current mode."""
        if self.input_mode.get() == self.SINGLE_MODE:
            return self.plasmid_qty.get()
        
        # For batch mode, find the key (identifier) that the folder name starts with
        for identifier, value in self.batch_replacements.items():
            if folder_name.startswith(identifier):
                return value
        return None # No match found

    def find_report_document(self, root_path: str) -> str:
        """
        Finds the most likely report document by walking a directory tree.
        This handles cases where the exact filename is not known.
        """
        all_word_docs = []

        # Walk through the directory to find all .doc and .docx files
        for dirpath, _, filenames in os.walk(root_path):
            for filename in filenames:
                # Ignore temporary Word files that start with '~'
                if (filename.lower().endswith(self.DOC_EXTENSION) or filename.lower().endswith(self.DOCX_EXTENSION)) and not filename.startswith(self.INVALID_FILE_PREFIX):
                    all_word_docs.append(os.path.join(dirpath, filename))

        # Case 1: No Word documents found
        if not all_word_docs:
            self.log(f"[ERROR] \t No Word documents found in {os.path.basename(root_path)}")
            return None

        # Case 2: Exactly one Word document found (ideal case)
        if len(all_word_docs) == 1:
            self.log(f"[OK] \t Found single Word document: {os.path.basename(all_word_docs[0])}")
            return all_word_docs[0]

        # Case 3: Multiple Word documents found. Apply heuristics to find the best match.
        self.log(f"[INFO] \t Multiple Word documents found in {os.path.basename(root_path)}. Selecting best match.")

        # Heuristic 1: Look for a specific substring in the filename.
        spaceless_substring = self.REPORT_SUBSTRING.replace(' ', '')
        candidates = [path for path in all_word_docs
                        if self.REPORT_SUBSTRING in os.path.basename(path).lower() or
                        spaceless_substring in os.path.basename(path).lower()]

        # If the substring match narrows it down to one file, use that.
        if len(candidates) == 1:
            self.log(f"[OK] \t Selected report by substring match: {os.path.basename(candidates[0])}")
            return candidates[0]

        # Heuristic 2 (Fallback): If there are still multiple candidates (or none from the substring search),
        # sort the list alphabetically and pick the first one.
        final_list = candidates if candidates else all_word_docs
        final_list.sort()
        chosen_doc_path = final_list[0]
        self.log(f"[WARNING] \t Multiple candidates or no substring match. Selecting first alphabetically: {os.path.basename(chosen_doc_path)}")
        return chosen_doc_path

    def find_target_table(self, doc_obj) -> Optional[Table]:
        """
        Iterates through all tables in a document to find the one containing the target text.
        
        Returns:
            The python-docx Table object if found, otherwise None.
        """
        for table in doc_obj.tables:
            for row in table.rows:
                # Check if the target text is in the first cell of the row
                if self.TARGET_CELL_TEXT in row.cells[self.TARGET_CELL_INDEX].text.lower():
                    return table
        return None

    def update_table_cell(self, table, new_qty: str, doc_name: str) -> bool:
        """
        Finds the target row within the given table and updates the quantity cell.
        
        Returns:
            True if the update was successful, False otherwise.
        """
        padded_qty = f"{self.QTY_PADDING}{new_qty}"

        # Iterate through the rows of the pre-identified table
        for row in table.rows:
            if self.TARGET_CELL_TEXT in row.cells[self.TARGET_CELL_INDEX].text.lower():
                try:
                    cell_to_update = row.cells[self.QTY_CELL_INDEX]
                    # Clear any existing content in the cell's first paragraph
                    p = cell_to_update.paragraphs[0]
                    p.clear()
                    run = p.add_run(padded_qty)
                    # Apply specific font formatting to match the document style
                    font = run.font
                    font.name = self.DEFAULT_FONT_NAME
                    font.size = Pt(self.DEFAULT_FONT_SIZE)
                    self.log(f"[OK] \t Updated 'Plasmid Qty' to '{new_qty}' in {doc_name}")
                    return True
                except IndexError:
                    self.log(f"[ERROR] \t Row for '{self.TARGET_CELL_TEXT}' in {doc_name} does not have enough cells to update.")
                    return False

        # This code path should ideally not be reached if the table was found correctly
        self.log(f"[WARNING] \t Could not find '{self.TARGET_CELL_TEXT}' row in the provided table for {doc_name}.")
        return False

    def process_word_document(self, doc_path: str, new_qty: str):
        """
        Main logic to open, modify, and save a single Word document.
        Handles both .doc and .docx formats by using win32com for conversion.
        """
        is_doc_format = doc_path.lower().endswith(self.DOC_EXTENSION)
        # Initialize variables
        word_app = None
        doc_name = os.path.basename(doc_path)
        temp_docx_path = None
        com_initialized = False

        try:
            doc_obj = None
            if is_doc_format:
                # Initialize COM for this thread to handle .doc files via Word application
                pythoncom.CoInitialize()
                com_initialized = True

                # For legacy .doc files, use COM to open Word in the background and
                # save the file as a temporary .docx, which python-docx can handle.
                # Use Dispatch instead of EnsureDispatch for better compatibility across Office versions.
                word_app = win32.Dispatch('Word.Application')
                word_app.Visible = False
                temp_docx_path = os.path.splitext(doc_path)[0] + self.TEMP_DOCX_SUFFIX
                
                doc = word_app.Documents.Open(doc_path)
                # Use SaveAs for older Word versions that may not have SaveAs2
                if hasattr(doc, 'SaveAs2'):
                    doc.SaveAs2(temp_docx_path, FileFormat=self.WIN32_DOCX_FORMAT)
                else:
                    doc.SaveAs(temp_docx_path, FileFormat=self.WIN32_DOCX_FORMAT)
                doc.Close(False)
                self.log(f"[OK] \t Converted {doc_name} to .docx for editing.")
                doc_obj = Document(temp_docx_path)
            else:
                # For modern .docx files, load them directly with python-docx
                doc_obj = Document(doc_path)

            # --- Edit the document object ---
            target_table = self.find_target_table(doc_obj)


            if target_table:
                self.update_table_cell(target_table, new_qty, doc_name)
            else:
                self.log(f"[WARNING] \t Could not find the target table in {doc_name}. Skipping update.")

            # --- Save changes ---
            if is_doc_format:
                # If we started with a .doc, save the modified temporary .docx,
                # then use COM again to save it back to the original .doc format.
                doc_obj.save(temp_docx_path)
                doc = word_app.Documents.Open(temp_docx_path)
                if hasattr(doc, 'SaveAs2'):
                    doc.SaveAs2(doc_path, FileFormat=self.WIN32_DOC_FORMAT)
                else:
                    doc.SaveAs(doc_path, FileFormat=self.WIN32_DOC_FORMAT)
                doc.Close(False)
                self.log(f"[OK] \t Saved changes back to original .doc format: {doc_name}")
            else:
                # Save the .docx file directly
                doc_obj.save(doc_path)
                self.log(f"[OK] \t Saved changes to .docx file: {doc_name}")

        except Exception as e:
            self.log(f"[ERROR] \t Failed during processing of {doc_name}: {e}")
        finally:
            # --- Cleanup ---
            # Ensure the Word application is closed, even if errors occurred.
            if word_app:
                word_app.Quit()
            # Delete the temporary .docx file if it was created.
            if temp_docx_path and os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)
            if com_initialized:
                # Uninitialize COM for this thread
                pythoncom.CoUninitialize()

    def zip_processed_folders(self, base_dir: str, zip_output_dir: str, folder_names: List[str]):
        """
        Zips each of the processed folders into a new archive in the output directory.
        """
        if not os.path.exists(zip_output_dir):
            os.makedirs(zip_output_dir)

        # Create a new .zip file for each processed folder
        for folder_name in folder_names:
            output_zip_path = os.path.join(zip_output_dir, folder_name)
            try:
                # To preserve the original zip structure (e.g., 'my_zip.zip' containing a 'my_zip' folder),
                # we need to set the 'root_dir' to the parent of the folder we are zipping,
                # and 'base_dir' to the name of the folder itself.
                shutil.make_archive(
                    base_name=output_zip_path,
                    format='zip',
                    root_dir=base_dir,
                    base_dir=folder_name
                )
                self.log(f"[OK] \t Zipped {folder_name} to {output_zip_path}.zip")
            except Exception as e:
                self.log(f"[ERROR] \t Failed to zip {folder_name}: {e}")
