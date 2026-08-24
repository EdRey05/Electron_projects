"""
Bio Basic Inc. - Canada
Made by: Eduardo Reyes, Ph.D.
Contact: ed5reyes@outlook

Version: 1.1
Date: Mar 30, 2026
Notes: Fixed issue with hub-passed directories.
"""

import os
import shutil
import zipfile
from typing import List, Tuple

import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import pythoncom
import win32com.client as win32

class InternalRenaming:
    def __init__(self, app_instance):
        """
        Initialize the workflow.
        - app_instance: A reference to the main App to access common elements.
        """
        self.app = app_instance
        self.log = self.app.log

        # --- Constants for this workflow ---
        self.PROCESSED_FOLDERS_DIR_NAME = "2-processed_folders_internal"
        self.OUTPUT_ZIPS_DIR_NAME = "3-output_zips_internal"
        self.DEFAULT_PATH_SUFFIX = "_internal"
        self.REPLACEMENT_FILE_DELIMITER = '\t'
        self.ZIP_EXTENSION = ".zip"
        self.TXT_EXTENSION = ".txt"
        self.DOC_EXTENSION = ".doc"
        self.DOCX_EXTENSION = ".docx"
        self.ZIP_FORMAT_NAME = "zip"
        self.REPORT_SUBSTRING = "gene synthesis report"
        self.QC_TABLE_PARAGRAPH_TEXT = "restriction enzyme digestion analysis"
        self.INFO_TABLE_INDEX = 0
        self.INFO_TABLE_ROW_INDEX = 0
        self.INFO_TABLE_CELL_INDEX = 1
        self.QC_TABLE_INDEX = 1
        self.DOCX_HEADER_FONT_SIZE = 16
        self.WIN32_DOC_FORMAT = 0
        self.WIN32_DOCX_FORMAT = 16
        self.TEMP_DOCX_SUFFIX = "_temp.docx"
        self.INVALID_FILE_PREFIX = '~'

        # --- Workflow-specific variables ---
        self.replacement_file = ctk.StringVar()
        self.replacement_map = {}
        self.copy_output_var = ctk.BooleanVar(value=False)
        self.delete_intermediate_var = ctk.BooleanVar(value=False)
        self.remove_qc_table_var = ctk.BooleanVar(value=False)

    def create_ui(self, parent_frame):
        """
        Creates the UI elements specific to this workflow inside the provided parent frame (a tab).
        """
        # --- Configure a 65/35 proportional layout for the tab panel ---
        parent_frame.grid_columnconfigure(0, weight=13) # 65%
        parent_frame.grid_columnconfigure(1, weight=7)  # 35%

        left_panel = ctk.CTkFrame(parent_frame, fg_color="transparent")
        left_panel.grid(row=0, column=0, sticky="nsew", padx=10)

        right_panel = ctk.CTkFrame(parent_frame, fg_color="transparent")
        right_panel.grid(row=0, column=1, sticky="nsew", padx=10)

        # Replacement map file selection
        replacement_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        replacement_frame.pack(pady=8, anchor="w")
        ctk.CTkLabel(replacement_frame, text="Replacement Map (.txt):").pack(side="left", padx=(0, 8))
        self.replacement_entry = ctk.CTkEntry(replacement_frame, textvariable=self.replacement_file, width=380)
        self.replacement_entry.pack(side="left", padx=(0, 8))
        ctk.CTkButton(replacement_frame, text="Select File", command=self.select_replacement_file).pack(side="left")

        # Optional step checkboxes
        ctk.CTkLabel(right_panel, text="Optional Steps:").pack(anchor="w", pady=(0, 4))
        ctk.CTkCheckBox(right_panel, text="Copy output to 'Received by BBI'", variable=self.copy_output_var).pack(anchor="w", pady=2)
        ctk.CTkCheckBox(right_panel, text="Delete intermediate files when finished", variable=self.delete_intermediate_var).pack(anchor="w", pady=2)
        ctk.CTkCheckBox(right_panel, text="Remove QC Table from Word docs", variable=self.remove_qc_table_var).pack(anchor="w", pady=2)

    def validate_inputs(self) -> bool:
        """
        Validates that all inputs required for this specific workflow are ready.
        """
        if not self.replacement_map:
            messagebox.showerror("Error", "Replacement map is not loaded. Please select a valid .txt file for this workflow.")
            return False

        if self.copy_output_var.get():
            default_output_path = self.app.make_path_from_config('Paths', 'QC_Output_Folder')
            if not default_output_path or not os.path.isdir(default_output_path):
                messagebox.showerror("Error",
                                    "The destination folder for copying output files ('Received by BBI') does not exist. "
                                    "Please check the path in config.ini or uncheck the copy option.")
                return False
        return True

    def run_processing_task(self):
        """
        The main processing logic for the Internal Renaming workflow.
        """
        self.log("[INFO] \t Starting Internal Renaming workflow...")
        input_folder = os.path.normpath(self.app.input_folder.get().strip())

        stats = {'processed': 0, 'skipped': 0, 'failed_rename': 0}
        all_zip_files = [f for f in os.listdir(input_folder) if f.lower().endswith(self.ZIP_EXTENSION)]
        total_zips = len(all_zip_files)
        zips_to_process = [f for f in all_zip_files if any(key in f for key in self.replacement_map.keys())]

        for zip_file in all_zip_files:
            if zip_file not in zips_to_process:
                self.log(f"[SKIP] \t Skipping {zip_file} as it does not match any replacement rule.")
                stats['skipped'] += 1

        stats['processed'] = len(zips_to_process)
        extracted_zips_dir, extracted_zips_names = self.extract_zip_files(input_folder, zips_to_process)
        if not extracted_zips_dir:
            self.log("[FINISHED] \t No zip files to process.")
            self.app.process_button.configure(state="normal")
            return

        final_folder_names = []
        for folder_name in extracted_zips_names:
            current_folder_path = os.path.join(extracted_zips_dir, folder_name)
            
            # Step 1: Rename all files and subfolders within the extracted directory
            self.rename_files_and_folders(current_folder_path)

            # Step 2: Find and process the single Word document within the directory
            self._find_and_process_report(current_folder_path)

            new_folder_name = folder_name
            for old, new in self.replacement_map.items():
                new_folder_name = new_folder_name.replace(old, new)

            if new_folder_name != folder_name:
                new_folder_path = os.path.join(extracted_zips_dir, new_folder_name)
                try:
                    shutil.move(current_folder_path, new_folder_path)
                    self.log(f"[OK] \t Renamed main folder: {folder_name} -> {new_folder_name}")
                    final_folder_names.append(new_folder_name)
                except Exception as e:
                    self.log(f"[ERROR] \t Could not rename main folder {folder_name}: {e}")
                    stats['failed_rename'] += 1
                    final_folder_names.append(folder_name)
            else:
                final_folder_names.append(folder_name)

        zip_output_dir = os.path.join(os.path.dirname(input_folder), self.OUTPUT_ZIPS_DIR_NAME)
        if final_folder_names:
            self.zip_processed_folders(extracted_zips_dir, zip_output_dir, final_folder_names)

        if self.copy_output_var.get():
            default_output_path = self.app.make_path_from_config('Paths', 'QC_Output_Folder')
            self.log("[OK] \t Copying output files to default location...")
            for zip_name in final_folder_names:
                source_zip_path = os.path.join(zip_output_dir, f"{zip_name}{self.ZIP_EXTENSION}")
                if os.path.exists(source_zip_path):
                    try:
                        shutil.copy(source_zip_path, default_output_path)
                        self.log(f"[OK] \t Copied {zip_name}.zip to {default_output_path}")
                    except Exception as e:
                        self.log(f"[ERROR] \t Failed to copy {zip_name}.zip: {e}")

        if self.delete_intermediate_var.get():
            self.log("[OK] \t Deleting intermediate files as requested...")
            for dir_path in [extracted_zips_dir, zip_output_dir]:
                if os.path.isdir(dir_path):
                    try:
                        shutil.rmtree(dir_path)
                        self.log(f"[OK] \t Deleted intermediate folder: {dir_path}")
                    except Exception as e:
                        self.log(f"[ERROR] \t Could not delete {dir_path}: {e}")

        self.log(f"[SUMMARY] \t Processed {stats['processed']}/{total_zips} relevant files ({stats['skipped']} skipped, {stats['failed_rename']} failed to rename).")
        self.log("[FINISHED] \t Renaming process complete!")
        self.app.process_button.configure(state="normal")

    # --- Helper Methods for this Workflow ---

    def setup_default_paths(self):
        default_path = self.app.make_path_from_config('Paths', 'QC_Input_Folder')
        if default_path:
            self.app.input_folder.set(default_path)
            self.log(f"[OK] \t Default input folder set to: {default_path}")

            txt_files = [f for f in os.listdir(default_path) if f.lower().endswith(self.TXT_EXTENSION)]
            if txt_files:
                map_path = os.path.join(default_path, txt_files[0])
                self.log(f"[OK] \t Auto-detected replacement map: {txt_files[0]}")
                self.replacement_file.set(map_path)
                self.load_replacement_strings()

    def select_replacement_file(self):
        file_selected = filedialog.askopenfilename(title="Select Replacement Map File", filetypes=[("Text files", f"*{self.TXT_EXTENSION}")])
        if file_selected:
            self.replacement_file.set(file_selected)
            self.load_replacement_strings()

    def load_replacement_strings(self):
        filepath = self.replacement_file.get()
        if not filepath or not os.path.isfile(filepath):
            self.log("[ERROR] \t Replacement map file is not valid.")
            return

        self.replacement_map.clear()
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                for i, line in enumerate(f):
                    if not line.strip(): continue
                    parts = line.rstrip('\n\r').split(self.REPLACEMENT_FILE_DELIMITER)
                    if len(parts) >= 1 and parts[0]:
                        old_str = parts[0]
                        new_str = parts[1] if len(parts) > 1 else ''
                        self.replacement_map[old_str] = new_str
                        self.log(f"[PLAN] \t Will replace '{old_str}' with '{new_str}'")
                    else:
                        self.log(f"[WARNING] \t Skipping malformed line {i+1} in replacement file.")
            self.log(f"[OK] \t Loaded {len(self.replacement_map)} replacement pairs.")
        except Exception as e:
            self.log(f"[ERROR] \t Failed to load replacement file: {e}")
            self.replacement_map.clear()

    def extract_zip_files(self, input_folder: str, zip_files_to_process: List[str]) -> Tuple[str, List[str]]:
        extracted_folders_dir = os.path.join(os.path.dirname(input_folder), self.PROCESSED_FOLDERS_DIR_NAME)
        if os.path.exists(extracted_folders_dir):
            shutil.rmtree(extracted_folders_dir)
        os.makedirs(extracted_folders_dir)
        self.log(f"[OK] \t Created clean processing folder: {extracted_folders_dir}")

        if not zip_files_to_process:
            return None, []

        for zip_file in zip_files_to_process:
            zip_path = os.path.join(input_folder, zip_file)
            try:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(extracted_folders_dir)
                    self.log(f"[OK] \t Extracted {zip_file}")
            except Exception as e:
                self.log(f"[ERROR] \t Failed to extract {zip_file}: {e}")

        return extracted_folders_dir, [os.path.splitext(f)[0] for f in zip_files_to_process]

    def rename_files_and_folders(self, root_path: str):
        for dirpath, dirnames, filenames in os.walk(root_path, topdown=False):
            for filename in filenames:
                new_filename = filename
                for old, new in self.replacement_map.items():
                    new_filename = new_filename.replace(old, new)

                old_filepath = os.path.join(dirpath, filename)
                new_filepath = os.path.join(dirpath, new_filename)

                if new_filename != filename and os.path.exists(old_filepath):
                    shutil.move(old_filepath, new_filepath)

            for dirname in dirnames:
                new_dirname = dirname
                for old, new in self.replacement_map.items():
                    new_dirname = new_dirname.replace(old, new)

                if new_dirname != dirname:
                    old_dirpath = os.path.join(dirpath, dirname)
                    new_dirpath = os.path.join(dirpath, new_dirname)
                    if os.path.exists(old_dirpath):
                        shutil.move(old_dirpath, new_dirpath)

    def _find_and_process_report(self, root_path: str):
        """Finds the single report document in the tree and processes it."""
        doc_path = self._find_report_document_in_tree(root_path)
        if doc_path:
            self.process_word_document(doc_path)

    def _find_report_document_in_tree(self, root_path: str) -> str:
        """Finds the most likely report document by walking a directory tree."""
        all_word_docs = []
        for dirpath, _, filenames in os.walk(root_path):
            for filename in filenames:
                if (filename.lower().endswith(self.DOC_EXTENSION) or filename.lower().endswith(self.DOCX_EXTENSION)) and not filename.startswith(self.INVALID_FILE_PREFIX):
                    all_word_docs.append(os.path.join(dirpath, filename))

        if not all_word_docs:
            self.log(f"[WARNING] \t No Word documents found in {os.path.basename(root_path)}")
            return None

        if len(all_word_docs) == 1:
            self.log(f"[OK] \t Found single Word document: {os.path.basename(all_word_docs[0])}")
            return all_word_docs[0]

        self.log(f"[INFO] \t Multiple Word documents found in {os.path.basename(root_path)}. Selecting best match.")
        spaceless_substring = self.REPORT_SUBSTRING.replace(' ', '')
        candidates = [path for path in all_word_docs
                        if self.REPORT_SUBSTRING in os.path.basename(path).lower() or
                        spaceless_substring in os.path.basename(path).lower()]

        if len(candidates) == 1:
            self.log(f"[OK] \t Selected report by substring match: {os.path.basename(candidates[0])}")
            return candidates[0]

        final_list = candidates if candidates else all_word_docs
        final_list.sort()
        chosen_doc_path = final_list[0]
        self.log(f"[WARNING] \t Multiple candidates or no substring match. Selecting first alphabetically: {os.path.basename(chosen_doc_path)}")
        return chosen_doc_path

    def process_word_document(self, doc_path: str):
        is_doc_format = doc_path.lower().endswith(self.DOC_EXTENSION)
        word_app = None
        temp_docx_path = None
        com_initialized = False

        try:
            doc_obj = None
            if is_doc_format:
                # Initialize COM for this thread to handle .doc files via Word application
                pythoncom.CoInitialize()
                com_initialized = True
                # Use Dispatch instead of EnsureDispatch for better compatibility across Office versions.
                # Dispatch is dynamic and less prone to issues with cached COM object information.
                word_app = win32.Dispatch('Word.Application')
                word_app.Visible = False
                temp_docx_path = os.path.splitext(doc_path)[0] + self.TEMP_DOCX_SUFFIX
                
                doc = word_app.Documents.Open(doc_path)
                # Use SaveAs for older Word versions that may not have SaveAs2
                if hasattr(doc, 'SaveAs2'):
                    doc.SaveAs2(temp_docx_path, FileFormat=self.WIN32_DOCX_FORMAT)
                else:
                    doc.SaveAs(temp_docx_path, FileFormat=self.WIN32_DOCX_FORMAT)
                doc.Close(False)
                doc_obj = Document(temp_docx_path)
            else:
                doc_obj = Document(doc_path)

            modified = False
            if len(doc_obj.tables) > 0:
                cell = doc_obj.tables[self.INFO_TABLE_INDEX].rows[self.INFO_TABLE_ROW_INDEX].cells[self.INFO_TABLE_CELL_INDEX]
                original_text = cell.text
                new_text = original_text
                for old, new in self.replacement_map.items():
                    new_text = new_text.replace(old, new)
                
                if new_text != original_text:
                    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    p.clear()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(new_text)
                    run.font.size = Pt(self.DOCX_HEADER_FONT_SIZE)
                    self.log(f"[OK] \t Updated text in Word doc: {os.path.basename(doc_path)}")
                    modified = True

            if self.remove_qc_table_var.get():
                para_to_delete = next((p for p in doc_obj.paragraphs if self.QC_TABLE_PARAGRAPH_TEXT in p.text.lower()), None)
                if para_to_delete:
                    p_element = para_to_delete._element
                    p_element.getparent().remove(p_element)
                    self.log(f"[OK] \t Removed QC header paragraph from {os.path.basename(doc_path)}")
                    modified = True

                if len(doc_obj.tables) > self.QC_TABLE_INDEX:
                    tbl_element = doc_obj.tables[self.QC_TABLE_INDEX]._element
                    tbl_element.getparent().remove(tbl_element)
                    self.log(f"[OK] \t Removed QC table from {os.path.basename(doc_path)}")
                    modified = True

            if modified:
                if is_doc_format:
                    doc_obj.save(temp_docx_path)
                    doc = word_app.Documents.Open(temp_docx_path)
                    if hasattr(doc, 'SaveAs2'):
                        doc.SaveAs2(doc_path, FileFormat=self.WIN32_DOC_FORMAT)
                    else:
                        doc.SaveAs(doc_path, FileFormat=self.WIN32_DOC_FORMAT)
                    doc.Close(False)
                else:
                    doc_obj.save(doc_path)
        except Exception as e:
            self.log(f"[ERROR] \t Failed to process Word document {os.path.basename(doc_path)}: {e}")
        finally:
            if temp_docx_path and os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)
            if word_app:
                word_app.Quit()
            if com_initialized:
                # Uninitialize COM for this thread
                pythoncom.CoUninitialize()

    def zip_processed_folders(self, base_dir: str, zip_output_dir: str, folder_names: List[str]):
        if not os.path.exists(zip_output_dir):
            os.makedirs(zip_output_dir)

        for folder_name in folder_names:
            output_zip_path = os.path.join(zip_output_dir, folder_name)
            try:
                # To preserve the original zip structure (e.g., 'my_zip.zip' containing a 'my_zip' folder),
                # we set 'root_dir' to the parent of the folder we are zipping,
                # and 'base_dir' to the name of the folder itself.
                shutil.make_archive(
                    base_name=output_zip_path,
                    format=self.ZIP_FORMAT_NAME,
                    root_dir=base_dir,
                    base_dir=folder_name
                )
                self.log(f"[OK] \t Zipped {folder_name} to {output_zip_path}{self.ZIP_EXTENSION}")
            except Exception as e:
                self.log(f"[ERROR] \t Failed to zip {folder_name}: {e}")
