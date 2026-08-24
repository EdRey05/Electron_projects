"""
Bio Basic Inc. - Canada
Made by: Eduardo Reyes, Ph.D.
Contact: ed5reyes@outlook

Version: 1.1
Date: Mar 30, 2026
Notes: Fixed issue with hub-passed directories.
"""

import customtkinter as ctk
from tkinter import messagebox
import os
import sys
import shutil
import zipfile
from typing import List, Optional, Tuple, Dict
from datetime import datetime

# Imports for Word processing
from docx import Document
from docx.table import Table
import win32com.client as win32
import pythoncom

class ResuspensionPrep:
    """
    A workflow for extracting QC data from two types of Word reports ("WGK" and "BATCH")
    found within zip archives. The extracted data is compiled into a single TSV text file.
    """
    def __init__(self, app_instance):
        """
        Initializes the ResuspensionPrep workflow.

        Args:
            app_instance: A reference to the main App to access common elements.
        """
        self.app = app_instance
        self.log = self.app.log

        # --- Constants for this workflow ---
        self.PROCESSED_FOLDERS_DIR_NAME = "2-processed_folders_resusp"
        self.DEFAULT_PATH_SUFFIX = "_resuspension"
        
        # File handling constants
        self.DOC_EXTENSION = ".doc"
        self.DOCX_EXTENSION = ".docx"
        self.INVALID_FILE_PREFIX = '~'
        self.REPORT_SUBSTRING = "gene synthesis report"
        self.TEMP_DOCX_SUFFIX = "_temp.docx"

        # Report type identifiers
        self.WGK_TYPE = "WGK"
        self.BATCH_TYPE = "BATCH"

        # --- Data storage ---
        self.extracted_data = []

    def create_ui(self, parent_frame):
        """
        Creates the user interface for the Resuspension Prep workflow.
        This UI is simpler as it does not require user input for values.
        """
        main_panel = ctk.CTkFrame(parent_frame, fg_color="transparent")
        main_panel.pack(fill="both", expand=True, padx=10, pady=10)

        info_text = "Extracts data into a Word table. The output is saved in the input folder."
        ctk.CTkLabel(main_panel, text=info_text, justify="center").pack(anchor="center", pady=(20, 10))

    def validate_inputs(self) -> bool:
        """
        Validates that the environment is suitable for running the workflow.
        """
        # This workflow uses win32com to handle .doc files, which is Windows-only
        if sys.platform != "win32":
            messagebox.showerror("Unsupported OS", "This workflow requires Windows to process .doc files.")
            return False
        
        return True

    def run_processing_task(self):
        """
        Orchestrates the entire data extraction process.
        """
        # Initialize COM for this thread
        if sys.platform == "win32":
            pythoncom.CoInitialize()
        
        try:
            self.log("[INFO] \t Starting Resuspension Prep workflow...")
            input_folder = os.path.normpath(self.app.input_folder.get().strip())
            all_zip_files = [f for f in os.listdir(input_folder) if f.lower().endswith('.zip')]

            if not all_zip_files:
                self.log("[WARNING] \t No zip files found in the input directory.")
                self.app.process_button.configure(state="normal")
                return

            # Extract all zip files to a temporary processing directory
            extracted_zips_dir, extracted_zips_names = self.extract_zip_files(input_folder, all_zip_files)
            if not extracted_zips_dir:
                self.log("[FINISHED] \t No zip files to process.")
                self.app.process_button.configure(state="normal")
                return

            # Prepare data list with the required header
            self.extracted_data = [["BBID", "Plasmid_ID", "Plasmid_Qty", "OD_260_280", "OD_260_230"]]
            
            processed_count = 0
            for folder_name in extracted_zips_names:
                current_folder_path = os.path.join(extracted_zips_dir, folder_name)
                
                # Determine the report type based on the folder name
                report_type = None
                if self.WGK_TYPE in folder_name.upper():
                    report_type = self.WGK_TYPE
                elif self.BATCH_TYPE in folder_name.upper():
                    report_type = self.BATCH_TYPE
                else:
                    self.log(f"[WARNING] \t Could not determine report type for '{folder_name}'. Skipping.")
                    continue

                # Find the Word document and process it to extract data
                doc_path = self.find_report_document(current_folder_path)
                if doc_path:
                    row_data = self.process_word_document(doc_path, current_folder_path, report_type)
                    if row_data:
                        # Append the extracted data in the correct order
                        self.extracted_data.append([
                            row_data.get("BBID", ""),
                            row_data.get("Plasmid_ID", ""),
                            row_data.get("Plasmid_Qty", ""),
                            row_data.get("OD_260_280", ""),
                            row_data.get("OD_260_230", "")
                        ])
                        processed_count += 1
            
            # Write the compiled data to a single output file
            self.write_output_file(input_folder)

            # Clean up intermediate files by default
            self.log("[OK] \t Deleting intermediate files...")
            if os.path.exists(extracted_zips_dir):
                try:
                    shutil.rmtree(extracted_zips_dir)
                except Exception as e:
                    self.log(f"[ERROR] \t Could not delete {extracted_zips_dir}: {e}")

            total_files = len(extracted_zips_names)
            self.log(f"[SUMMARY] \t Extracted data from {processed_count}/{total_files} reports.")
            self.log("[FINISHED] \t Resuspension Prep process complete!")
        except Exception as e:
            self.log(f"[FATAL] \t An unexpected error occurred in the main task: {e}")
        finally:
            # Ensure the UI button is re-enabled
            self.app.process_button.configure(state="normal")
            # Uninitialize COM for this thread
            if sys.platform == "win32":
                pythoncom.CoUninitialize()

    def write_output_file(self, output_dir: str) -> Optional[str]:
        """
        Writes the collected data to a table in a new Word (.docx) file.
        """
        if len(self.extracted_data) <= 1:
            self.log("[INFO] \t No data was extracted, so no output file will be created.")
            return None
        
        # Ensure the output directory exists
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Create a unique filename using a timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"Resuspension_Prep_Data_{timestamp}.docx"
        output_filepath = os.path.join(output_dir, output_filename)

        try:
            # Create a new Word document
            doc = Document()
            doc.add_heading('Resuspension Prep Data', level=1)

            # Add a table with rows and columns based on the extracted data
            table = doc.add_table(rows=len(self.extracted_data), cols=len(self.extracted_data[0]))
            table.style = 'Table Grid'

            # Populate the table with headers and data
            for i, row_data in enumerate(self.extracted_data):
                for j, cell_data in enumerate(row_data):
                    cell = table.cell(i, j)
                    cell.text = str(cell_data)
                    # Make the header row bold
                    if i == 0:
                        cell.paragraphs[0].runs[0].bold = True
            
            # Adjust column widths to fit content
            table.autofit = True
            
            # Save the document
            doc.save(output_filepath)
            self.log(f"[OK] \t Successfully wrote data to {output_filename}")
            return output_filepath
        except Exception as e:
            self.log(f"[ERROR] \t Failed to write output file: {e}")
            return None

    def process_word_document(self, doc_path: str, folder_path: str, report_type: str) -> Optional[Dict[str, str]]:
        """
        Main logic to open, read, and extract data from a single Word document.
        Handles both .doc and .docx formats by using win32com for conversion.
        """
        is_doc_format = doc_path.lower().endswith(self.DOC_EXTENSION)
        word = None
        doc_name = os.path.basename(doc_path)
        temp_docx_path = None
        data = {}

        try:
            doc_obj = None
            if is_doc_format:
                # For legacy .doc files, convert to a temporary .docx for reading
                # Use Dispatch instead of EnsureDispatch for better compatibility across Office versions.
                word = win32.Dispatch('Word.Application')
                word.Visible = False
                temp_docx_path = os.path.splitext(doc_path)[0] + self.TEMP_DOCX_SUFFIX
                
                doc = word.Documents.Open(doc_path)
                # Use SaveAs for older Word versions that may not have SaveAs2
                if hasattr(doc, 'SaveAs2'):
                    doc.SaveAs2(temp_docx_path, FileFormat=16) # wdFormatXMLDocument
                else:
                    doc.SaveAs(temp_docx_path, FileFormat=16) # wdFormatXMLDocument
                doc.Close(False)
                self.log(f"[OK] \t Converted {doc_name} to .docx for reading.")
                doc_obj = Document(temp_docx_path)
            else:
                # For modern .docx files, load them directly
                doc_obj = Document(doc_path)

            # --- Delegate data extraction based on report type ---
            if report_type == self.BATCH_TYPE:
                data = self.extract_data_from_batch_doc(doc_obj)
            elif report_type == self.WGK_TYPE:
                data = self.extract_data_from_wgk_doc(doc_obj)
                # WGK reports require an additional step to get the Plasmid ID
                data["Plasmid_ID"] = self.extract_plasmid_id_from_wgk_folder(folder_path)

            self.log(f"[OK] \t Extracted data from {doc_name}")
            return data

        except Exception as e:
            self.log(f"[ERROR] \t Failed during processing of {doc_name}: {e}")
            return None
        finally:
            # --- Cleanup ---
            if word:
                word.Quit()
            if temp_docx_path and os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)

    def extract_data_from_batch_doc(self, doc_obj: Document) -> Dict[str, str]:
        """
        Extracts data from a 'BATCH' type report by searching its tables.
        This function is optimized to stop searching as soon as data is found.
        """
        data = {"OD_260_230": ""} # This value is not present in BATCH reports
        
        # Table 1: BBID, Plasmid No., Plasmid Qty.
        if len(doc_obj.tables) > 0:
            found_count = 0
            table1 = doc_obj.tables[0]
            for row in table1.rows:
                for i, cell in enumerate(row.cells):
                    # Search for labels and grab the text from the cell to the right
                    cell_text = cell.text.strip().lower()
                    if "bbi id" in cell_text and i + 1 < len(row.cells):
                        data["BBID"] = row.cells[i+1].text.strip()
                        found_count += 1
                    if "plasmid no." in cell_text and i + 1 < len(row.cells):
                        data["Plasmid_ID"] = row.cells[i+1].text.strip()
                        found_count += 1
                    if "plasmid qty." in cell_text and i + 1 < len(row.cells):
                        data["Plasmid_Qty"] = row.cells[i+1].text.strip()
                        found_count += 1
                
                # If all 3 items are found, we can stop searching this table
                if found_count >= 3:
                    break
                    
        # Table 3: OD260/OD280
        if len(doc_obj.tables) > 2:
            table3 = doc_obj.tables[2]
            for r, row in enumerate(table3.rows):
                for c, cell in enumerate(row.cells):
                    if "od260/od280" in cell.text.lower().replace(" ", ""):
                        # Value is in the cell directly below the label
                        if r + 1 < len(table3.rows):
                            data["OD_260_280"] = table3.cell(r + 1, c).text.strip()
                        return data # Exit once found
        return data

    def extract_data_from_wgk_doc(self, doc_obj: Document) -> Dict[str, str]:
        """
        Extracts data from a 'WGK' type report by searching its tables.
        This function is optimized to stop searching as soon as data is found.
        """
        data = {}

        # Table 1: BBI ID
        if len(doc_obj.tables) > 0:
            found_bbi_id = False
            table1 = doc_obj.tables[0]
            for row in table1.rows:
                for i, cell in enumerate(row.cells):
                    if "bbi id" in cell.text.lower() and i + 2 < len(row.cells):
                        # For WGK, the value is two cells to the right of the label
                        data["BBID"] = row.cells[i+2].text.strip()
                        found_bbi_id = True
                        break # Exit inner cell loop
                if found_bbi_id:
                    break # Exit outer row loop

        # Table 2: Plasmid Qty
        if len(doc_obj.tables) > 1:
            found_qty = False
            table2 = doc_obj.tables[1]
            for row in table2.rows:
                for i, cell in enumerate(row.cells):
                    if "plasmid qty" in cell.text.lower() and i + 2 < len(row.cells):
                        data["Plasmid_Qty"] = row.cells[i+2].text.strip()
                        found_qty = True
                        break
                if found_qty:
                    break

        # Table 3: OD ratios
        if len(doc_obj.tables) > 2:
            found_280 = "OD_260_280" in data
            found_230 = "OD_260_230" in data
            table3 = doc_obj.tables[2]
            for row in table3.rows:
                for i, cell in enumerate(row.cells):
                    cell_text_lower = cell.text.lower()
                    if not found_280 and "od: 260/280" in cell_text_lower and i + 2 < len(row.cells):
                        data["OD_260_280"] = row.cells[i+2].text.strip()
                        found_280 = True
                    if not found_230 and "od: 260/230" in cell_text_lower and i + 2 < len(row.cells):
                        data["OD_260_230"] = row.cells[i+2].text.strip()
                        found_230 = True
                # If both are found, stop searching this table
                if found_280 and found_230:
                    break
        return data

    def extract_plasmid_id_from_wgk_folder(self, folder_path: str) -> str:
        """
        Finds the Plasmid ID from an .ab1 file in a subfolder for WGK reports.
        """
        seq_results_folder = os.path.join(folder_path, "OriginalSequencingResults")
        # Check if the specific subfolder exists
        if not os.path.isdir(seq_results_folder):
            self.log(f"[WARNING] \t 'OriginalSequencingResults' folder not found in {os.path.basename(folder_path)}")
            return ""
        
        for filename in os.listdir(seq_results_folder):
            if filename.lower().endswith(".ab1"):
                # The Plasmid ID is the part of the filename before the first dot
                plasmid_id = filename.split('.')[0]
                self.log(f"[OK] \t Found Plasmid ID '{plasmid_id}' from file '{filename}'")
                return plasmid_id
        
        self.log(f"[WARNING] \t No .ab1 files found in 'OriginalSequencingResults' for {os.path.basename(folder_path)}")
        return ""

    def setup_default_paths(self):
        """
        Sets up the default input folder for this workflow in the main application UI.
        """
        # Get the base path for input folders from the shared app configuration
        default_path = self.app.make_path_from_config('Paths', 'QC_Input_Folder')
        if default_path and os.path.isdir(default_path):
            self.app.input_folder.set(default_path)
            self.log(f"[OK] \t Default input folder found and set to: {default_path}")

    def extract_zip_files(self, input_folder: str, zip_files: List[str]) -> Tuple[str, List[str]]:
        """
        Extracts a list of zip files into a clean temporary directory for processing.
        """
        # Define the path for the temporary directory where files will be extracted
        extracted_dir = os.path.join(os.path.dirname(input_folder), self.PROCESSED_FOLDERS_DIR_NAME)
        
        # Ensure the extraction directory is clean by deleting it if it exists
        if os.path.exists(extracted_dir):
            shutil.rmtree(extracted_dir)
        # Recreate the empty directory
        os.makedirs(extracted_dir)

        # Loop through each zip file and extract its contents
        for zip_file in zip_files:
            zip_path = os.path.join(input_folder, zip_file)
            try:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(extracted_dir)
                self.log(f"[OK] \t Extracted {zip_file}")
            except Exception as e:
                self.log(f"[ERROR] \t Failed to extract {zip_file}: {e}")

        # Return the path to the extraction directory and a list of the folder names
        return extracted_dir, [os.path.splitext(f)[0] for f in zip_files]

    def find_report_document(self, root_path: str) -> str:
        """
        Finds the most likely report document within a directory by applying heuristics.
        This method is robust against inconsistent file naming.
        """
        all_word_docs = []
        # Walk through the directory to find all .doc and .docx files
        for dirpath, _, filenames in os.walk(root_path):
            for filename in filenames:
                # Ignore temporary Word files that often start with '~'
                if (filename.lower().endswith(self.DOC_EXTENSION) or filename.lower().endswith(self.DOCX_EXTENSION)) and not filename.startswith(self.INVALID_FILE_PREFIX):
                    all_word_docs.append(os.path.join(dirpath, filename))

        # Case 1: No Word documents found.
        if not all_word_docs:
            self.log(f"[ERROR] \t No Word documents found in {os.path.basename(root_path)}")
            return None

        # Case 2: Exactly one Word document found (the ideal scenario).
        if len(all_word_docs) == 1:
            self.log(f"[OK] \t Found single Word document: {os.path.basename(all_word_docs[0])}")
            return all_word_docs[0]

        # Case 3: Multiple Word documents found. Apply heuristics to find the best match.
        self.log(f"[INFO] \t Multiple Word documents found in {os.path.basename(root_path)}. Selecting best match.")
        # Heuristic 1: Look for a specific substring in the filename.
        spaceless_substring = self.REPORT_SUBSTRING.replace(' ', '')
        candidates = [path for path in all_word_docs
                        if self.REPORT_SUBSTRING in os.path.basename(path).lower() or
                        spaceless_substring in os.path.basename(path).lower()]

        # If the substring search narrows it down to a single file, use that one.
        if len(candidates) == 1:
            self.log(f"[OK] \t Selected report by substring match: {os.path.basename(candidates[0])}")
            return candidates[0]

        # Heuristic 2 (Fallback): If heuristics fail, sort alphabetically and pick the first one.
        final_list = candidates if candidates else all_word_docs
        final_list.sort()
        chosen_doc_path = final_list[0]
        self.log(f"[WARNING] \t Multiple candidates or no substring match. Selecting first alphabetically: {os.path.basename(chosen_doc_path)}")
        return chosen_doc_path